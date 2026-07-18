"""Fill the Appraisal Sheet and Report of Rating docx templates."""
import copy
import io
import os
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxcompose.composer import Composer

def resource_path(rel):
    """Resolve a bundled resource, both in dev and inside a PyInstaller exe."""
    base = getattr(sys, "_MEIPASS", os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base, rel)


APPRAISAL_TEMPLATES = {
    "BSCS": resource_path(os.path.join("templates_docx", "Appraisal_Sheet_BSCS.docx")),
    "BSIT": resource_path(os.path.join("templates_docx", "Appraisal_Sheet_BSIT.docx")),
    "BLIS": resource_path(os.path.join("templates_docx", "Appraisal_Sheet_BLIS.docx")),
    "DCT": resource_path(os.path.join("templates_docx", "Appraisal_Sheet_DCT.docx")),
}
DEFAULT_APPRAISAL_COURSE = "BSCS"
REPORT_TPL = resource_path(os.path.join("templates_docx", "Report_of_Rating.docx"))

# Appraisal Sheet year-block headings, in document order. Each maps to the
# term tables that follow it (2 for a full year, 1 for Mid Year), used to
# drop whole unused blocks (e.g. Third Year/Mid Year/Fourth Year for a
# student who only reached Second Year).
YEAR_GROUP_LABELS = ("FIRST YEAR", "SECOND YEAR", "THIRD YEAR",
                     "MID YEAR", "FOURTH YEAR")

FONT = "Calibri"


def _set_cell_valign_top(cell):
    """Anchor cell content to the top (Word sometimes inherits bottom
    alignment from the table style, pushing values to the lower line)."""
    from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    v = tc_pr.find(qn('w:vAlign'))
    if v is None:
        v = tc_pr.makeelement(qn('w:vAlign'), {})
        tc_pr.append(v)
    v.set(qn('w:val'), 'top')


def _set_cell(cell, text, size=10, center=False, bold=False):
    # keep exactly one paragraph in the cell
    for extra in cell.paragraphs[1:]:
        extra._p.getparent().remove(extra._p)
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    run = p.add_run(str(text))
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cell_valign_top(cell)


def _total_units(subjects):
    total = 0.0
    for s in subjects:
        try:
            total += float(s["units"])
        except ValueError:
            pass
    return f"{total:.1f}"


def _add_data_row(table, template_row):
    """Clone a data row so tables can grow past the template's blank rows."""
    new = copy.deepcopy(template_row._tr)
    template_row._tr.addnext(new)


def _delete_row(table, row):
    row._tr.getparent().remove(row._tr)


def _fill_table_rows(table, subjects, first_data, set_row, trim=True):
    """Grow a table's data rows to fit `subjects` (cloning the template's
    last data row as needed), fill each with `set_row(row, subject)`, and
    optionally delete any rows left blank between the last subject and the
    trailing TOTAL row."""
    last_data = len(table.rows) - 2
    n_slots = last_data - first_data + 1
    while n_slots < len(subjects):
        _add_data_row(table, table.rows[last_data])
        last_data += 1
        n_slots += 1
    for i, subj in enumerate(subjects):
        set_row(table.rows[first_data + i], subj)
    if trim:
        for idx in range(last_data, first_data + len(subjects) - 1, -1):
            _delete_row(table, table.rows[idx])


W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def _remove_unused_year_groups(doc, n_used_slots):
    """Delete whole Appraisal Sheet year-blocks (heading + its term tables)
    that fall entirely beyond the terms actually being filled. Must run
    after every table has been filled, since it only looks at how many
    term-table slots precede each heading."""
    body = doc.element.body
    children = list(body.iterchildren())
    p_tag, tbl_tag = f'{W_NS}p', f'{W_NS}tbl'

    def text_of(el):
        # Headings live inside floating text boxes, whose itertext() also
        # picks up numeric shape-position values and duplicated fallback
        # content, so match by substring rather than exact equality.
        return "".join(el.itertext())

    headings = []
    for i, el in enumerate(children):
        if el.tag != p_tag:
            continue
        txt = text_of(el)
        if any(label in txt for label in YEAR_GROUP_LABELS):
            headings.append(i)
    if not headings:
        return
    end_idx = next((i for i, el in enumerate(children)
                    if el.tag == p_tag and "Appraised by:" in text_of(el)),
                   len(children))

    slot = 0
    for gi, start in enumerate(headings):
        stop = headings[gi + 1] if gi + 1 < len(headings) else end_idx
        span = children[start:stop]
        n_tables = sum(1 for el in span if el.tag == tbl_tag)
        if slot >= n_used_slots:
            for el in span:
                body.remove(el)
        slot += n_tables


def _strip_trailing_empty_paragraphs(doc):
    """Remove empty paragraphs at the end of the body (they can create a
    blank trailing page, like in the Report of Rating template)."""
    body = doc.element.body
    for el in reversed(list(body.iterchildren())):
        tag = el.tag.split('}')[1]
        if tag == 'sectPr':
            continue
        is_empty_p = (tag == 'p'
                      and not ''.join(el.itertext()).strip()
                      and not el.findall(f'.//{W_NS}drawing'))
        if is_empty_p:
            body.remove(el)
        else:
            break


def term_sy_label(term, sy):
    short = term.replace(" Term", "").strip()          # "1st"
    yy = sy.split("-")                                  # 2025-2026 -> 25-26
    return f"{short}/{yy[0][2:]}-{yy[1][2:]}" if len(yy) == 2 else f"{short}/{sy}"


# ---------------------------------------------------------------- appraisal

def appraisal_template_for(course):
    return APPRAISAL_TEMPLATES.get((course or "").strip().upper(),
                                   APPRAISAL_TEMPLATES[DEFAULT_APPRAISAL_COURSE])


def build_appraisal(student, faculty_map=None, adviser="", dean=""):
    """Return a filled Document for one student (all terms). The template
    (BSCS/BSIT/BLIS/DCT) is picked from the student's course, detected by
    the parser from the grade listing PDF."""
    faculty_map = faculty_map or {}
    doc = Document(appraisal_template_for(student.get("course")))

    # Name / ID line
    for p in doc.paragraphs:
        if "Name" in p.text and "ID Number" in p.text:
            for r in list(p.runs):
                r._r.getparent().remove(r._r)
            for text, bold in [("Name: ", True), (student["name"], False),
                               ("\t\t", False), ("ID Number: ", True),
                               (student["id"], False)]:
                run = p.add_run(text)
                run.bold = bold
                run.font.name = FONT
                run.font.size = Pt(11)
            break

    # sort the student's terms: earlier SY first, 1st before 2nd
    order = {"1st Term": 0, "2nd Term": 1, "3rd Term": 2, "4th Term": 3}
    keys = sorted(student["terms"].keys(), key=lambda k: (k[1], order.get(k[0], 9)))

    n_slots = len(doc.tables) - 2  # last two tables are the signature blocks
    for slot, key in enumerate(keys[:n_slots]):
        term, sy = key
        subjects = student["terms"][key]
        table = doc.tables[slot]
        label = term_sy_label(term, sy)

        def set_row(row, subj, label=label):
            _set_cell(row.cells[0], subj["code"], size=8, center=True)
            _set_cell(row.cells[1], subj["title"], size=8)
            _set_cell(row.cells[2], subj["units"], size=8, center=True)
            _set_cell(row.cells[3], subj["grade"], size=8, center=True)
            _set_cell(row.cells[4],
                      faculty_map.get(subj["code"].replace(" ", ""), ""), size=8)
            _set_cell(row.cells[5], label, size=8, center=True)

        _fill_table_rows(table, subjects, 1, set_row, trim=True)
        totals = table.rows[-1]
        _set_cell(totals.cells[2], _total_units(subjects), size=8, center=True,
                  bold=True)

    n_used = min(len(keys), n_slots)
    # A year group can be kept (its first table has data) while its second
    # table has none, e.g. a student who just started Third Year - that
    # trailing table was never touched by the loop above, so it's still
    # the full blank template; trim it down to just header + totals too.
    for slot in range(n_used, n_slots):
        table = doc.tables[slot]
        _fill_table_rows(table, [], 1, lambda row, subj: None, trim=True)
        _set_cell(table.rows[-1].cells[2], _total_units([]), size=8, center=True,
                  bold=True)

    _remove_unused_year_groups(doc, n_used)

    if adviser:
        _set_cell(doc.tables[-2].rows[0].cells[0], adviser, size=11, bold=True)
    if dean:
        _set_cell(doc.tables[-1].rows[0].cells[0], dean, size=11, bold=True)

    _strip_trailing_empty_paragraphs(doc)
    return doc


def fill_appraisal(student, faculty_map=None, adviser="", dean=""):
    return _to_buffer(build_appraisal(student, faculty_map, adviser, dean))


# ----------------------------------------------------------------- report

def build_report(student, term, sy, subjects, faculty_map=None, adviser="",
                 dean="", trim_rows=True):
    """Return a filled Document for one student and one term."""
    faculty_map = faculty_map or {}
    doc = Document(REPORT_TPL)

    info = doc.tables[0]
    _set_cell(info.rows[0].cells[1], student["name"], size=11, bold=True)
    section = subjects[0]["section"] if subjects else student.get("course", "")
    _set_cell(info.rows[0].cells[3], section, size=11, bold=True)
    _set_cell(info.rows[1].cells[1], term, size=11, bold=True)
    _set_cell(info.rows[1].cells[3], sy, size=11, bold=True)

    def set_row(row, subj):
        _set_cell(row.cells[0], subj["code"], center=True)
        _set_cell(row.cells[1], subj["title"])
        _set_cell(row.cells[2], subj["units"], center=True)
        _set_cell(row.cells[3], subj["rating"], center=True)
        _set_cell(row.cells[4], subj["grade"], center=True)
        _set_cell(row.cells[5], faculty_map.get(subj["code"].replace(" ", ""), ""))

    grades = doc.tables[1]
    _fill_table_rows(grades, subjects, 2, set_row, trim=trim_rows)
    totals = grades.rows[-1]
    _set_cell(totals.cells[2], _total_units(subjects), center=True, bold=True)

    sig = doc.tables[2]
    if adviser:
        _set_cell(sig.rows[0].cells[0], adviser, size=11, bold=True)
    if dean:
        _set_cell(sig.rows[0].cells[2], dean, size=11, bold=True)

    _strip_trailing_empty_paragraphs(doc)
    return doc


def fill_report(student, term, sy, subjects, faculty_map=None, adviser="",
                dean="", trim_rows=True):
    return _to_buffer(build_report(student, term, sy, subjects,
                                   faculty_map, adviser, dean, trim_rows))


# ----------------------------------------------------------------- combine

def _page_break_before(doc):
    """Mark the first paragraph of a document to start on a new page."""
    from docx.oxml.ns import qn
    body = doc.element.body
    first_p = None
    for el in body.iterchildren():
        if el.tag == f'{W_NS}p':
            first_p = el
            break
        if el.tag == f'{W_NS}tbl':
            # body starts with a table: insert a tiny paragraph before it
            p = body.makeelement(qn('w:p'), {})
            el.addprevious(p)
            first_p = p
            break
    if first_p is None:
        return
    pPr = first_p.find(qn('w:pPr'))
    if pPr is None:
        pPr = first_p.makeelement(qn('w:pPr'), {})
        first_p.insert(0, pPr)
    if pPr.find(qn('w:pageBreakBefore')) is None:
        pPr.append(pPr.makeelement(qn('w:pageBreakBefore'), {}))


def combine_documents(docs):
    """Merge several filled Documents into one, each starting on a new page."""
    if not docs:
        return None
    composer = Composer(docs[0])
    for doc in docs[1:]:
        _page_break_before(doc)
        composer.append(doc)
    return _to_buffer(composer.doc)


def _to_buffer(doc):
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
